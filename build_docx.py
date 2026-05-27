import docx
from docx.shared import Pt, Inches
import re

def markdown_to_docx(md_path, docx_path):
    doc = docx.Document()
    
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    blocks = content.split('\n\n')
    
    for block in blocks:
        block = block.strip()
        if not block:
            continue
            
        if block == '---' or block.startswith('<div'):
            continue
            
        # Check for headings
        if block.startswith('# '):
            doc.add_heading(block[2:].strip(), 0)
        elif block.startswith('## '):
            doc.add_heading(block[3:].strip(), 1)
        elif block.startswith('### '):
            doc.add_heading(block[4:].strip(), 2)
        elif block.startswith('<img'):
            # Extract src
            match = re.search(r'src="([^"]+)"', block)
            if match:
                img_path = '/Users/sandilyachimalamarri/Plateplanner/' + match.group(1)
                try:
                    doc.add_picture(img_path, width=Inches(3.5))
                except Exception as e:
                    print(f"Skipping image {img_path}: {e}")
        elif block.startswith('- ') or block.startswith('* '):
            lines = block.split('\n')
            for line in lines:
                if line.startswith('- ') or line.startswith('* '):
                    doc.add_paragraph(line[2:].strip(), style='List Bullet')
                elif line.strip():
                     doc.add_paragraph(line.strip())
        elif re.match(r'^\d+\.\s', block):
            lines = block.split('\n')
            for line in lines:
                if re.match(r'^\d+\.\s', line):
                    match = re.match(r'^\d+\.\s(.*)', line)
                    if match:
                        doc.add_paragraph(match.group(1).strip(), style='List Number')
                elif line.strip():
                     doc.add_paragraph(line.strip())
        else:
            # Clean up bolding markup a bit
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', block)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)

    doc.save(docx_path)

if __name__ == '__main__':
    markdown_to_docx('/Users/sandilyachimalamarri/Plateplanner/CMPE_295B_Report_Filled.md', '/Users/sandilyachimalamarri/Plateplanner/PlatePlanner_295B_Report.docx')
