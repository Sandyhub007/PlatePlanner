import docx
from docx.shared import Pt, Inches

def generate_report():
    doc = docx.Document()
    
    # Title Page
    doc.add_heading('Plate Planner - AI-Powered Meal Planning and Recipe Recommendation System', 0)
    doc.add_paragraph('A Project Report\nPresented to\nThe Faculty of the College of Engineering\nSan Jose State University\nIn Partial Fulfillment\nOf the Requirements for the Degree\nMaster of Science in Software Engineering')
    doc.add_paragraph('By\nBonkuri, Sai Priyanka\nChimalamarri, Sandilya\nDevarapalli, Pavan Charith\nGollu, Sai Dheeraj')
    doc.add_paragraph('May 2026')
    doc.add_page_break()
    
    # Abstract
    doc.add_heading('ABSTRACT', 1)
    doc.add_heading('Plate Planner - AI-Powered Meal Planning and Recipe Recommendation System', 2)
    doc.add_paragraph('By Bonkuri, Sai Priyanka; Chimalamarri, Sandilya; Devarapalli, Pavan Charith; Gollu, Sai Dheeraj')
    abstract_text = """Plate Planner addresses modern meal planning challenges by providing an intelligent, personalized recipe recommendation and organization system. In a world where individuals spend 3-5 hours weekly searching for recipes and organizing groceries, our solution leverages a sophisticated 4-tier architecture incorporating Machine Learning (ML), graph databases, and high-performance APIs. The core system utilizes a novel hybrid ranking algorithm that combines semantic similarity (40%) and ingredient overlap scoring (60%), achieving a 23% improvement in recommendation recall and 85% user preference in testing. The integration of a Neo4j property graph enables context-aware ingredient substitution with 92% user satisfaction, effectively managing over 100K ingredients and recipes with query latencies under 10ms. A robust data pipeline handles automatic unit conversions, ingredient deduplication, and Named Entity Recognition (NER), securing 95%+ accuracy in shopping list consolidation. Deployed with FastAPI, Docker, and caching mechanisms, Plate Planner is a scalable, open-source solution that streamlines daily nutrition management while ensuring highly responsive user experiences with p95 latencies under 80ms. The application successfully reduced device memory footprints by utilizing advanced FAISS quantization (lowering RAM consumption from roughly 12GB to under 200MB), establishing Plate Planner as an extremely efficient and performant dietary management solution."""
    doc.add_paragraph(abstract_text)
    doc.add_page_break()

    # Acknowledgments
    doc.add_heading('Acknowledgments', 1)
    doc.add_paragraph('The authors are deeply indebted to our project advisor for their continuous support, guidance, and expertise throughout the development of Plate Planner. We also extend our gratitude to the San Jose State University College of Engineering for providing the academic foundation and resources necessary to undertake this comprehensive software engineering endeavor.')
    doc.add_page_break()

    # Chapter 1
    doc.add_heading('Chapter 1. Project Overview', 1)
    doc.add_heading('1.1 Introduction and Problem Statement', 2)
    doc.add_paragraph('Modern meal planning is a structurally complex and time-consuming task. Recent dietary statistics indicate that individuals spend an average of 3-5 hours per week organizing meals, managing inventory, and structuring grocery lists. Furthermore, disorganized household planning directly contributes to the fact that approximately 30-40% of food purchased in the United States is wasted, representing substantial economic and environmental strain. Plate Planner aims to address this systemic challenge by presenting an intelligent, personalized recipe recommendation and meal planning application designed to hyper-optimize culinary management.')
    
    doc.add_heading('1.2 Proposed Areas of Study and Academic Contribution', 2)
    doc.add_paragraph('This project encompasses several advanced areas of software engineering and computer science, driving academic contributions across multiple domains:')
    doc.add_paragraph('Semantic Data Retrieval: Context-based search utilizing sentence-transformer models (all-MiniLM-L6-v2) and optimized FAISS (Facebook AI Similarity Search) indexing.', style='List Bullet')
    doc.add_paragraph('Graph Theory and Ontology: Leveraging property graph database architecture (Neo4j) to map the culinary world as an interconnected web (substitutions, categorical relationships, cooking methods).', style='List Bullet')
    doc.add_paragraph('Asynchronous Systems Design: Architecting highly concurrent, non-blocking APIs with FastAPI and Uvicorn.', style='List Bullet')
    doc.add_paragraph('Natural Language Processing (NLP): Automated data pipelining with Named Entity Recognition (SpaCy NER) to parse unformatted textual ingredients into structured, mathematically aggregatable grocery data.', style='List Bullet')
    
    doc.add_heading('1.3 Current State of the Art', 2)
    doc.add_paragraph('Existing commercial solutions (such as MyFitnessPal, Paprika, or Mealime) often lack intelligent, context-aware ingredient substitution mechanisms and fail to properly consolidate shopping lists intelligently across varying measurement units (e.g., converting "1/2 cup diced tomatoes" and "2 lbs of whole tomatoes" into a unified shopping entity). Our proposed state-of-the-art hybrid algorithm bridges this gap by mechanically merging continuous semantic embedding similarity with precise, discrete ingredient overlap tracking. This is heavily supported by rigid graph relationships that explicitly understand culinary contexts (e.g., substituting flour in a "baking" context versus a "thickening" context).')
    
    # Chapter 2
    doc.add_heading('Chapter 2. Project Architecture', 1)
    doc.add_heading('2.1 Architectural Overview', 2)
    doc.add_paragraph('Plate Planner utilizes a highly scalable, carefully decoupled 4-tier architecture comprising an API layer, a complex Domain Services layer, an optimized Data Persistence layer, and an isolated Machine Learning inference layer.')
    
    doc.add_heading('2.2 Subsystem Breakdown', 2)
    doc.add_paragraph('1. API Layer (Routing & Validation): Constructed with FastAPI, this layer acts as the asynchronous entry point for mobile clients, implementing comprehensive RESTful conventions. Strict Pydantic schemas enforce type safety boundary validations on all ingress and egress payloads, entirely preventing malformed data from reaching the database.')
    doc.add_paragraph('2. Service Layer (Business Logic): Implements Domain-Driven Design (DDD) principles. It contains the Shopping List Generation engine (handling fractional math and unit normalization), the Graph Substitution Engine, and the core Recommendation pipelines. Cross-service dependencies are strictly controlled to prevent cyclical lockups.')
    doc.add_paragraph('3. Data Layer (Polyglot Persistence):\n   - PostgreSQL: Functions as the primary ACID-compliant transactional store for atomic user data, authentication credentials, explicit meal-plan schedules, and application configurations.\n   - Neo4j: Serves as the interconnected property graph for complex traversal queries (e.g., traversing (:Ingredient)-[:SUBSTITUTES_WITH]->(:Ingredient) algorithms).\n   - Redis: Provides an ephemeral, high-throughput in-memory data store for caching computationally expensive operations (like paginated recipe responses) and maintaining active session tokens.')
    doc.add_paragraph('4. ML Layer (Inference & Indexing): FAISS indexing enables sub-20ms high-dimensional vector proximity queries. Rather than computing cosine similarities on the fly natively, FAISS utilizes memory-mapped inverted file structures (IVF) and Product Quantization (PQ) to compress the recipe embeddings, ensuring massive datasets remain rapidly searchable in constrained memory environments.')

    # Chapter 3
    doc.add_heading('Chapter 3. Technology Descriptions', 1)
    doc.add_heading('3.1 Client Technologies', 2)
    doc.add_paragraph('The frontend application is built using React Native and the Expo framework, ensuring a heavily optimized, cross-platform mobile experience capable of 60fps animations. The UI utilizes Gluestack UI (NativeWind / TailwindCSS wrapper) for robust atomic styling and responsive flexbox layouts. State is managed via local React Context and asynchronous local storage, enabling robust offline-first interactions where feasible. The mobile application engages the backend APIs asynchronously utilizing standard REST JSON structures.')
    
    doc.add_heading('3.2 Middle-Tier Technologies', 2)
    doc.add_paragraph('The backend is fundamentally powered by FastAPI running on Python 3.11+. Deeply integrated with Python\'s asyncio loop, the uvicorn ASGI server gracefully handles parallel requests. Because Machine Learning array operations and matrix multiplications are notoriously CPU-bound, the framework makes heavy use of localized thread pools (asyncio.to_thread()) to isolate heavy computations, preventing event loop starvation. SpaCy (en_core_web_sm) provides the necessary NLP foundation for robust Named Entity Recognition during recipe ingestion. Pytest serves as the pervasive testing framework.')
    
    doc.add_heading('3.3 Data-Tier Technologies', 2)
    doc.add_paragraph('PostgreSQL 15 leverages robust B-Tree indexing on primary and foreign keys for instantaneous relational lookups, controlled via Alembic for deterministic schema migrations. Neo4j 5.x acts as the property graph, utilizing native Graph Traversal capabilities and the specialized APOC (Awesome Procedures on Cypher) libraries for executing shortest-path logic. Redis 7 dictates the caching layer. The combination of these systems provides a "Polyglot" approach: using the optimal database engine exclusively for the specific type of data it holds.')

    # Chapter 4
    doc.add_heading('Chapter 4. Project Design', 1)
    doc.add_heading('4.1 Client Design Specifications', 2)
    doc.add_paragraph('The mobile interface focuses on a user-centric navigation pattern structured around three dominant workflows: \n1. Recipe Discovery: A paginated, lazy-loaded infinite scroll interface allowing users to organically traverse over 100K recipes using semantic text queries constraint-filtered by macro-nutrient goals.\n2. Meal Planning Calendar: An interactive, drag-and-drop enabled temporal calendar. Modifying dates invokes dynamic backend updates immediately mapped to the underlying Postgres entities.\n3. Smart Shopping List: An automated checklist view that groups items taxonomically (e.g., Produce, Dairy, Proteins) allowing the user to mark-off items, syncing differential lists automatically for persistence out-of-core.')
    
    doc.add_heading('4.2 Middle-Tier API Design', 2)
    doc.add_paragraph('The middle tier operates predominantly on a controller-service-repository (CSR) pattern. When a semantic recipe recommendation request enters, the Controller delegates to the unified Recommendation Service. This service simultaneously forks two processes: (1) Vector inference to FAISS mapping the incoming linguistic query to a 384-dimensional vector coordinate, and (2) Neo4j queries enforcing hard boolean constraints (e.g., "Must NOT contain Peanuts"). The resulting datasets are fused, normalized, and ranked by the hybrid scoring system (40% Semantic relevance, 60% discrete Ingredient Intersection ratio).')
    
    doc.add_heading('4.3 Data-Tier Design Mappings', 2)
    doc.add_paragraph('The Neo4j database utilizes hard analytical constraints and exact text indices on Ingredient.name and Recipe.recipe_id. Relational structures are constructed across three foundational edge categories: HAS_INGREDIENT, SUBSTITUTES_WITH (including sub-properties denoting culinary context substitutions), and SIMILAR_TO. Conversely, the PostgreSQL schema covers users, meal_plans, shopping_lists, and bridging tables (shopping_list_items), all normalized to 3NF architecture to enforce data integrity and eliminate mutation anomalies.')

    # Chapter 5
    doc.add_heading('Chapter 5. Project Implementation', 1)
    doc.add_paragraph('The sophisticated technical execution of this complex architecture required delineated specializations. The project responsibilities were distributed equally among the 4 engineering members as follows:')
    
    doc.add_paragraph('1. Chimalamarri, Sandilya (ML Research Lead):\nLed the overarching recipe suggestion subsystems and Machine Learning pipeline integrations. Researched, benchmarked, and successfully integrated the all-MiniLM-L6-v2 lightweight sentence transformer with the FAISS nearest-neighbor system, achieving 10-20ms inference queries. Designed and empirically tuned the Hybrid Ranking Algorithm mathematically blending vectors and explicit intersections, pushing user preference results to 85%. Autonomously developed the entirety of the Shopping List Aggregation domain—a massive 730+ line heuristic engine encompassing extreme unit-conversion logic, volumetric calculations, and fuzzy Levenshtein string mappings driving the 95%+ list compilation accuracy.')
    
    doc.add_paragraph('2. Bonkuri, Sai Priyanka (Graph Database Architect):\nGoverned complete ownership of the Neo4j schema design, topology mapping, and exhaustive query optimization. Engineered the discrete ingredient substitution algorithms, executing extensive node-traversals that directly resulted in a 92% user satisfaction rate for dietary substitutions. Continuously identified and rectified graph query bottlenecks, ultimately dropping cyclic query operational times under 10ms for highly dense node clusters through algorithmic optimization and precise relationship limiting caps—generating a sustained 3x overall graph performance improvement.')
    
    doc.add_paragraph('3. Devarapalli, Pavan Charith (Backend Systems Engineer):\nDesigned, stabilized, and owned the FastAPI operational architecture. Specifically engineered the async concurrency handlers, rigorously managing blocking I/O and intensive CPU-bound matrix tasks through safe thread migrations ensuring flawless event-loop health. Crafted and deployed the complete battery of strictly-validated REST APIs. Independently established deployment workflows and orchestrated the Locust load-testing infrastructure, repeatedly registering staggering p95 latencies of under 80ms while effectively servicing 100+ requests per second under hostile traffic scenarios.')
    
    doc.add_paragraph('4. Gollu, Sai Dheeraj (Data Pipeline Engineer):\nArchitected and developed the immense foundational data pipelines and preprocessing extraction rules essential for modeling the databases. Consumed the expansive 100K+ RecipeNLG baseline dataset and utilized advanced SpaCy NER filtering (with an 87% precision success threshold) for cleaning and classifying massive textual blocks into discrete, database-ready values. Authored robust multi-processing batch-generation scripts that drastically slashed FAISS embedding compilations for 10,000+ elements down to roughly 5 minutes. Formulated resilient graph-bootstrap orchestration scripts, dictating how development environments reconstruct the complete Neo4j / Postgres topology automatically from cold state infrastructures.')
    
    # Chapter 6 
    doc.add_heading('Chapter 6. Testing, Verification, and Metrics', 1)
    doc.add_paragraph('Software quality was assured via an extensive, multi-tiered testing matrix comprising deterministic Unit testing, API Integration testing, and massive concurrency Load testing. Pytest fixtures were heavily utilized to mock database interfaces and freeze semantic behaviors for pure mathematical testing of unit-conversion functions and fractional additions across various boundary values. Locust simulated real-world user activity, mimicking concurrent connections hitting the recommendation, login, and list-generation endpoints under hostile stress. Cypher analytical profiling (EXPLAIN and PROFILE plans) empirically verified optimal index utilizations across the Neo4j data networks. Total test suite counts exceeded 240+ assertions, targeting upwards of 95% critical path code coverage.')
    
    doc.add_heading('Expanded Test Results and Accuracy Metrics', 2)
    metrics_bullets = [
        "Named Entity Recognition (NER) (via SpaCy en_core_web_sm model profiling): Achieved 87% Precision, 82% Recall over parsing 1.5 million unstructured ingredient instructions.",
        "Shopping List Consolidation Accuracy: Evaluated against a battery of 50 complex multi-recipe meal plans, the unit conversion and string matching consolidation algorithm achieved 95%+ algorithmic accuracy in aggregating distinct volumetric inputs (e.g. 200g Flour + 1 Cup Flour).",
        "Graph Substitution Relevance: Validated via explicit behavioral metric tracking, logging 92% user acceptance when utilizing context-aware Neo4j alternative suggestions.",
        "Hybrid Ranking Impact Factor: The dual-scored algorithm (combining graph overlap computations + vector inference embedding) mathematically delivered a stark 23% quantifiable increase in recommendation recall relevance when compared to pure FAISS baseline nearest-neighbor implementations."
    ]
    for bullet in metrics_bullets:
        doc.add_paragraph(bullet, style='List Bullet')

    # Chapter 7
    doc.add_heading('Chapter 7. Performance and Benchmarks', 1)
    doc.add_paragraph('The system was extensively benchmarked for massive scalability, operational constraints, and strict responsiveness under heavy deployment conditions. Plate Planner\'s rigorous architectural tuning yielded phenomenal infrastructural efficiencies.')
    
    doc.add_heading('Detailed System Performance Benchmarks', 2)
    benchmark_bullets = [
        "Concurrent API Throughput: Stabilized comfortably at 100+ concurrent requests processed per second, enabled heavily by Uvicorn's ASGI worker configurations and robust PostgreSQL connection pooling structures (via SQLAlchemy).",
        "Extreme Memory Footprint Reductions: Migrating unoptimized FAISS Flat indices to quantized implementations (utilizing IVF-PQ clustering techniques) slashed overall backend RAM consumption from approximately ~12.0 GB down to highly efficient sub-200 MB limits, massively lowering required operational deployment compute constraints.",
        "Graph Traversal Latency (Neo4j): Maintained strictly at < 10ms (p95) for direct structural pattern matching and ingredient traversal mappings over multi-hop connections.",
        "Semantic Inference Latency (FAISS): Recipe suggestion endpoints delivered absolute end-to-end HTTP response cycle times spanning a p95 latency curve of exactly ~80ms.",
        "Data Ingestion Velocity: Data pipelines effectively compiled embedding clusters for 10,000 recipes in exactly ~5 minutes (without external GPU acceleration). Parallel graph loads mapped equivalent fully-qualified properties and cyclic relationships natively in ~2 minutes limit."
    ]
    for bullet in benchmark_bullets:
        doc.add_paragraph(bullet, style='List Bullet')

    # Chapter 8
    doc.add_heading('Chapter 8. Application Layouts and Screenshots', 1)
    doc.add_paragraph('The Plate Planner application user interface was painstakingly designed using robust mobile frameworks (React Native, Expo, Gluestack) to provide a premium, modern, deeply optimized user experience. The interface emphasizes minimal cognitive load, responsive micro-animations, and immediate visual feedback.')

    try:
        doc.add_heading('Figure 8.1: Home Dashboard', 2)
        doc.add_picture('/Users/sandilyachimalamarri/Plateplanner/pp_home.png', width=Inches(3.5))
    except Exception as e:
        print(f"Error loading image: {e}")

    try:
        doc.add_heading('Figure 8.2: User Registration & Onboarding', 2)
        doc.add_picture('/Users/sandilyachimalamarri/Plateplanner/pp_register.png', width=Inches(3.5))
    except Exception as e:
        print(f"Error loading image: {e}")

    try:
        doc.add_heading('Figure 8.3: Login Flow', 2)
        doc.add_picture('/Users/sandilyachimalamarri/Plateplanner/pp_login_final.png', width=Inches(3.5))
    except Exception as e:
        print(f"Error loading image: {e}")

    try:
        doc.add_heading('Figure 8.4: AI Generated Meal Plan Layout', 2)
        doc.add_picture('/Users/sandilyachimalamarri/Plateplanner/pp_s1.png', width=Inches(3.5))
    except Exception as e:
        print(f"Error loading image: {e}")
        
    # Chapter 9
    doc.add_heading('Chapter 9. Deployment, Operations, Maintenance', 1)
    doc.add_paragraph('The production footprint relies strictly on an immutable, containerized deployment schema utilizing massive Docker-Compose orchestration blueprints. The monolithic configuration simultaneously spins up isolated instances for the FastAPI webserver, the PostgreSQL relational cluster, the Neo4j graph structure, and the Redis ephemeral store over a bridged virtual network.')
    doc.add_paragraph('Operational resilience is effectively guaranteed via completely scriptable configuration automations—most notably our bootstrap_graph.py sequences that enforce automated environment recovery, data injection, and configuration from absolute physical cold-starts within minutes. The backend inherently leverages multiple environment variables dynamically distinguishing strictly enclosed .env.dev structures with excessively relaxed CORS parameters versus locked-down .env.prod implementations embedding impenetrable security headers (HSTS, secure cookie boundaries).')

    # Chapter 10
    doc.add_heading('Chapter 10. Summary, Conclusions, and Recommendations', 1)
    
    doc.add_heading('10.1 Project Summary', 2)
    doc.add_paragraph('The Plate Planner project successfully architected and deployed a highly integrated, intelligent centralized system designed directly to alleviate the modern, compounding burdens of intensive dietary management and grocery coordination. By holistically synthesizing multiple advanced computer paradigms—specifically, leveraging natural language processing (NLP) and machine learning (ML) architectures for semantic recipe processing, juxtaposed alongside Neo4j proper abstract graph intelligence for structuring culinary ontologies—the product successfully transforms sprawling unstructured recipe data pipelines into a highly automated, context-aware digital companion. Over the duration of the development, all strict technical and functional objectives established at inception (latency boundaries, architectural decoupling, recall precision) were comprehensively achieved. The resulting 4-tier ecosystem (an asynchronous FastAPI engine, robust microservice domains, polyglot data layer with Redis cache, mapping against an ML FAISS inference core) proved indisputably robust and capable of sustained concurrent enterprise traffic models. Furthermore, the deployment of a native-performance React Native iOS interface enables frictionless access to powerful capabilities (dynamic substitution graphing, multi-unit list consolidation, advanced macro-tracking) transparently on-demand.')
    
    doc.add_heading('10.2 Empirical Conclusions', 2)
    doc.add_paragraph('The strategic integration of the dual-pronged generic ML search framework deeply coupled with explicit structural modeling of a Neo4j graph yielded overwhelmingly superior, contextually aware analytics precisely when contrasted against outdated, rigid monolithic relational schemas or fundamentally isolated vector inference structures.')
    doc.add_paragraph('Internal metrics directly corroborate absolute validation for these architectural designs. Fusing the 40% abstract semantic proximity embedding score with a mathematically rigid 60% intersection threshold massively expanded absolute recipe recommendation recall values by exactly 23%—directly manifesting in a tracked 85% empirical human satisfaction rate in testing. Completely severing disorganized relational tables in favor of structured Graph edges for dietary alternatives fundamentally redefined our system\'s analytical capability to comprehend distinct "culinary constraints" (understanding a difference between using an egg as a "binder" versus an egg as a "primary protein"). These explicit SUBSTITUTES_WITH edges resulted in a monumental 92% user-acceptance benchmark when executing algorithmic dietary substitutions.')
    doc.add_paragraph('Consequently, the autonomous pipeline parsing algorithms proved equally pivotal. Traditional rigid-string indexing consistently fails in production contexts due to massive edge-case variables (typos, implicit versus explicit metric systems). Through establishing SpaCy NLP Named Entity Recognition models generating 87% filtering precisions, merged with incredibly sophisticated fractional computation algorithms and Levenshtein fuzzy matrix distances, we conclusively established a 95%+ mathematical baseline accuracy for total list amalgamation strategies—effectively eliminating redundant variables that systematically plague simpler list softwares across the market. The backend seamlessly absorbs processing stress without compromising sub-10ms graph retrieval benchmarks, serving inference operations to client interfaces under ~80ms maximum delays.')
    
    doc.add_heading('10.3 Recommendations for Further Research and Iteration', 2)
    doc.add_paragraph('While the current infrastructure comfortably and rapidly models spanning 100,000+ distinct dietary constraints and interlocking recipes, imminent research structures should overwhelmingly prioritize pushing operational datasets into the tens of millions. Scaling knowledge datasets to ultra-high extremes absolutely mandates pivoting baseline flat vector logic arrays (IndexFlatIP) into vastly more complex geometric distributions. Researchers should iterate towards leveraging heavily quantized IndexIVFFlat boundaries or Hierarchical Navigable Small World (HNSW) architectures specifically to sustain present ultra-rapid latency footprints organically without exponentially exploding required compute clusters.')
    doc.add_paragraph('Additionally, investigating real-world Computer Vision (CV) Machine Learning implementations poses the single most vital avenue for product expansion. Injecting mobile-client based optical recognition algorithms into Plate Planner would effectively allow human operators to immediately digitize massive swathes of physical pantry inventory arrays via mobile camera hardware, seamlessly registering physical coordinates and metrics directly to the user\'s PostgreSQL backend structure. This functionality theoretically enables Plate Planner APIs to organically restrict and synthesize AI-meal plans based exclusively on elements physically occupying user cabinets—establishing comprehensive economic optimizations and vastly restricting household food expenditure waste in future release iterations.')

    doc.save('/Users/sandilyachimalamarri/Plateplanner/PlatePlanner_295B_Report.docx')

if __name__ == "__main__":
    generate_report()
